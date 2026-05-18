"""DOCX extractor emits markdown source on ``body_md`` for the new
structural preview renderer.

Bold and italic runs round-trip as ``**…**`` / ``*…*``; bulleted and
numbered list-styled paragraphs become GFM list lines; tables become
GFM pipe tables. ``body_struct`` continues to carry plain-text Block
records so snippets don't show literal markdown markers.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from fnd.extract.docx import extract


@pytest.fixture
def docx_path(tmp_path: Path) -> Path:
    """Build a docx fixture covering bold/italic runs, bulleted list,
    numbered list, and a 2x2 table — all under one Heading 1."""
    doc = Document()
    doc.add_heading("Top Section", level=1)
    p = doc.add_paragraph()
    p.add_run("plain ")
    p.add_run("bold").bold = True
    p.add_run(" and ")
    p.add_run("italic").italic = True
    doc.add_paragraph("first bullet", style="List Bullet")
    doc.add_paragraph("second bullet", style="List Bullet")
    doc.add_paragraph("first numbered", style="List Number")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "alpha"
    table.cell(1, 1).text = "beanstalk"
    out = tmp_path / "fixture.docx"
    doc.save(str(out))
    return out


def test_body_md_carries_bold_and_italic(docx_path: Path) -> None:
    chunk = next(extract(docx_path))
    md = chunk.body_md
    assert "**bold**" in md, md
    assert "*italic*" in md, md


def test_body_md_carries_bulleted_list(docx_path: Path) -> None:
    md = next(extract(docx_path)).body_md
    assert "- first bullet" in md
    assert "- second bullet" in md


def test_body_md_carries_numbered_list(docx_path: Path) -> None:
    md = next(extract(docx_path)).body_md
    assert "1. first numbered" in md


def test_body_md_carries_pipe_table(docx_path: Path) -> None:
    md = next(extract(docx_path)).body_md
    assert "| Item | Score |" in md
    assert "| alpha | beanstalk |" in md
    # GFM separator row is present.
    assert "|------|" in md.replace(" ", "")


def test_body_md_starts_with_heading(docx_path: Path) -> None:
    md = next(extract(docx_path)).body_md
    assert md.lstrip().startswith("# Top Section")


def test_body_struct_remains_plain_text(docx_path: Path) -> None:
    """Snippet pipeline reads body_struct — markdown markers in body_md
    must not leak into block.text."""
    chunk = next(extract(docx_path))
    for block in chunk.body_struct:
        assert "**" not in block.text, block.text
        assert "*" not in block.text or block.text.startswith("*") is False, block.text


def test_table_cells_are_indexed(docx_path: Path) -> None:
    """Cell text should be searchable — landed in chunk.body so F_BODY
    sees it."""
    chunk = next(extract(docx_path))
    assert "beanstalk" in chunk.body, chunk.body


def test_paragraph_table_paragraph_interleaves_in_md(tmp_path: Path) -> None:
    """A table sandwiched between two paragraphs lands in its real
    document position in body_md (not relegated to the end)."""
    doc = Document()
    doc.add_heading("Mix", level=1)
    doc.add_paragraph("alpha paragraph before table")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "TABLEMARK"
    doc.add_paragraph("beta paragraph after table")
    p = tmp_path / "mix.docx"
    doc.save(str(p))
    md = next(extract(p)).body_md
    a = md.index("alpha paragraph before table")
    t = md.index("TABLEMARK")
    b = md.index("beta paragraph after table")
    assert a < t < b, md


def test_no_heading_produces_single_chunk(tmp_path: Path) -> None:
    """A docx with no Heading-styled paragraphs still flushes one chunk
    when body content exists."""
    doc = Document()
    doc.add_paragraph("Just one paragraph, no heading style.")
    p = tmp_path / "no_heading.docx"
    doc.save(str(p))
    chunks = list(extract(p))
    assert len(chunks) == 1
    assert "Just one paragraph" in chunks[0].body_md


def test_heading_path_unchanged_with_subheadings(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("Top", level=1)
    doc.add_paragraph("intro")
    doc.add_heading("Mid", level=2)
    doc.add_paragraph("middle")
    doc.add_heading("Leaf", level=3)
    doc.add_paragraph("leaf body")
    p = tmp_path / "nested.docx"
    doc.save(str(p))
    chunks = list(extract(p))
    paths = [c.heading_path for c in chunks]
    assert paths == ["Top", "Top > Mid", "Top > Mid > Leaf"]
